# -*- coding: utf-8 -*-
"""
Data Analyzer - 数据分析核心模块
支持多格式文件分析、统计、可视化
版本: 1.0.0
"""

import os
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# 可选依赖
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # 无头模式
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False


class DataAnalyzer:
    """数据分析器主类"""
    
    SUPPORTED_FORMATS = ['.xlsx', '.xls', '.csv', '.docx', '.pdf', '.txt', '.md', '.markdown']
    
    def __init__(self, folder_path: str = None):
        self.folder = Path(folder_path) if folder_path else None
        self.files = {'excel': [], 'csv': [], 'word': [], 'pdf': [], 'txt': [], 'markdown': []}
        self.analysis_results = []
        
        if self.folder and self.folder.exists():
            self._scan_files()
    
    def _scan_files(self):
        """扫描文件夹中的支持的文件"""
        if not self.folder:
            return
            
        for f in self.folder.rglob('*'):
            if f.is_file():
                ext = f.suffix.lower()
                if ext in ['.xlsx', '.xls']:
                    self.files['excel'].append(str(f))
                elif ext == '.csv':
                    self.files['csv'].append(str(f))
                elif ext == '.docx':
                    self.files['word'].append(str(f))
                elif ext == '.pdf':
                    self.files['pdf'].append(str(f))
                elif ext == '.txt':
                    self.files['txt'].append(str(f))
                elif ext in ['.md', '.markdown']:
                    self.files['markdown'].append(str(f))
    
    def analyze_excel(self, file_path: str) -> Dict[str, Any]:
        """分析Excel文件"""
        try:
            df = pd.read_excel(file_path)
            return {
                'success': True,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
                'numeric_columns': list(df.select_dtypes(include=['number']).columns),
                'missing_values': df.isnull().sum().to_dict(),
                'basic_stats': df.describe().to_dict() if not df.empty else {},
                'data': df
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_csv(self, file_path: str, encoding: str = 'utf-8') -> Dict[str, Any]:
        """分析CSV文件"""
        try:
            # 尝试不同编码
            for enc in [encoding, 'gbk', 'gb2312', 'latin1']:
                try:
                    df = pd.read_csv(file_path, encoding=enc)
                    break
                except:
                    continue
            
            return {
                'success': True,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()},
                'numeric_columns': list(df.select_dtypes(include=['number']).columns),
                'missing_values': df.isnull().sum().to_dict(),
                'basic_stats': df.describe().to_dict() if not df.empty else {},
                'data': df
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_word(self, file_path: str) -> Dict[str, Any]:
        """分析Word文件"""
        if not HAS_DOCX:
            return {'success': False, 'error': 'python-docx未安装'}
        
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'success': True,
                'paragraphs': len(paragraphs),
                'total_text': '\n'.join(paragraphs),
                'tables_count': len(tables),
                'tables': tables
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_pdf(self, file_path: str) -> Dict[str, Any]:
        """分析PDF文件"""
        if not HAS_PYMUPDF:
            return {'success': False, 'error': 'PyMuPDF未安装'}
        
        try:
            doc = fitz.open(file_path)
            text = ''
            for page in doc:
                text += page.get_text()
            
            result = {
                'success': True,
                'pages': len(doc),
                'text_length': len(text),
                'text_preview': text[:500] if text else ''
            }
            doc.close()
            return result
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_txt(self, file_path: str, encoding: str = 'utf-8') -> Dict[str, Any]:
        """分析文本文件"""
        try:
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            lines = content.split('\n')
            words = content.split()
            
            return {
                'success': True,
                'lines': len(lines),
                'words': len(words),
                'characters': len(content),
                'preview': content[:500] if content else ''
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def analyze_markdown(self, file_path: str) -> Dict[str, Any]:
        """分析Markdown文件"""
        return self.analyze_txt(file_path)  # 复用txt分析
    
    def analyze_file(self, file_path: str, encoding: str = 'utf-8') -> Dict[str, Any]:
        """自动检测并分析任何支持的文件"""
        ext = Path(file_path).suffix.lower()
        
        analyzers = {
            'excel': self.analyze_excel,
            'csv': self.analyze_csv,
            'word': self.analyze_word,
            'pdf': self.analyze_pdf,
            'txt': self.analyze_txt,
            'markdown': self.analyze_markdown
        }
        
        for ftype, analyzer in analyzers.items():
            if ext in getattr(self.files, 'excel', ['.xlsx', '.xls']):
                return self.analyze_excel(file_path)
            elif ext == '.csv':
                return self.analyze_csv(file_path, encoding)
            elif ext == '.docx':
                return self.analyze_word(file_path)
            elif ext == '.pdf':
                return self.analyze_pdf(file_path)
            elif ext == '.txt':
                return self.analyze_txt(file_path, encoding)
            elif ext in ['.md', '.markdown']:
                return self.analyze_markdown(file_path)
        
        return {'success': False, 'error': f'不支持的格式: {ext}'}
    
    def generate_summary(self) -> Dict[str, Any]:
        """生成分析摘要"""
        summary = {
            'total_files': 0,
            'file_count_by_type': {},
            'file_details': [],
            'errors': []
        }
        
        for ftype, flist in self.files.items():
            summary['file_count_by_type'][ftype] = len(flist)
            summary['total_files'] += len(flist)
            
            for fpath in flist:
                try:
                    analysis = self.analyze_file(fpath)
                    self.analysis_results.append({
                        'name': os.path.basename(fpath),
                        'path': fpath,
                        'type': ftype,
                        'analysis': analysis
                    })
                    summary['file_details'].append({
                        'name': os.path.basename(fpath),
                        'type': ftype,
                        'status': 'success'
                    })
                except Exception as e:
                    summary['errors'].append({
                        'file': fpath,
                        'error': str(e)
                    })
        
        return summary
    
    def compare_files(self, file_paths: List[str]) -> Dict[str, Any]:
        """对比多个文件"""
        results = []
        for fpath in file_paths:
            result = self.analyze_file(fpath)
            results.append(result)
        
        return {
            'files_count': len(results),
            'results': results
        }
    
    def merge_data(self, file_paths: List[str], output_path: str = None) -> pd.DataFrame:
        """合并多个数据文件"""
        dfs = []
        
        for fpath in file_paths:
            ext = Path(fpath).suffix.lower()
            if ext in ['.xlsx', '.xls', '.csv']:
                try:
                    if ext == '.csv':
                        df = pd.read_csv(fpath, encoding='utf-8', errors='ignore')
                    else:
                        df = pd.read_excel(fpath)
                    df['_source_file'] = os.path.basename(fpath)
                    dfs.append(df)
                except:
                    continue
        
        if dfs:
            merged = pd.concat(dfs, ignore_index=True)
            if output_path:
                ext = Path(output_path).suffix.lower()
                if ext == '.csv':
                    merged.to_csv(output_path, index=False, encoding='utf-8-sig')
                else:
                    merged.to_excel(output_path, index=False)
            return merged
        
        return pd.DataFrame()
    
    def generate_statistics(self, df: pd.DataFrame) -> Dict[str, Any]:
        """生成数据统计"""
        stats = {
            'row_count': len(df),
            'column_count': len(df.columns),
            'numeric_columns': [],
            'categorical_columns': [],
            'missing_data': {}
        }
        
        if df.empty:
            return stats
        
        for col in df.columns:
            if df[col].dtype in ['int64', 'float64']:
                stats['numeric_columns'].append({
                    'name': col,
                    'mean': float(df[col].mean()),
                    'median': float(df[col].median()),
                    'std': float(df[col].std()),
                    'min': float(df[col].min()),
                    'max': float(df[col].max())
                })
            else:
                stats['categorical_columns'].append({
                    'name': col,
                    'unique_count': int(df[col].nunique()),
                    'top_values': df[col].value_counts().head(5).to_dict()
                })
        
        stats['missing_data'] = df.isnull().sum().to_dict()
        
        return stats
    
    def export_report(self, format: str = 'markdown', output_path: str = None, 
                     title: str = '数据分析报告') -> str:
        """导出分析报告"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        if format == 'markdown':
            content = f"""# {title}

**生成时间：** {timestamp}

## 摘要统计

- **总文件数：** {self.summary.get('total_files', 0)}
{self._format_file_count()}

## 详细分析

"""
            for detail in self.analysis_results:
                content += f"""### {detail['name']}

- **类型：** {detail['type']}
- **状态：** {'成功' if detail['analysis'].get('success') else '失败'}

"""
                if detail['analysis'].get('success'):
                    analysis = detail['analysis']
                    if 'rows' in analysis:
                        content += f"- **行数：** {analysis['rows']}\n"
                        content += f"- **列数：** {analysis['columns']}\n"
                    if 'pages' in analysis:
                        content += f"- **页数：** {analysis['pages']}\n"
                    if 'text_length' in analysis:
                        content += f"- **文字长度：** {analysis['text_length']}\n"
                
                content += "\n"
            
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            return content
        
        elif format == 'excel':
            if not output_path:
                output_path = 'analysis_report.xlsx'
            
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                summary_df = pd.DataFrame([{
                    '文件名': d['name'],
                    '类型': d['type'],
                    '状态': '成功' if d['analysis'].get('success') else '失败'
                } for d in self.analysis_results])
                summary_df.to_excel(writer, sheet_name='摘要', index=False)
                
                # 写入详细数据
                for detail in self.analysis_results:
                    if detail['analysis'].get('success') and 'data' in detail['analysis']:
                        df = detail['analysis']['data']
                        sheet_name = detail['name'][:31]  # Excel限制31字符
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            return f"报告已保存到: {output_path}"
        
        return "不支持的格式"
    
    def _format_file_count(self) -> str:
        """格式化文件数量"""
        lines = []
        for ftype, count in self.summary.get('file_count_by_type', {}).items():
            if count > 0:
                emoji = {
                    'excel': '📊',
                    'csv': '📄',
                    'word': '📝',
                    'pdf': '📕',
                    'txt': '📃',
                    'markdown': '📋'
                }.get(ftype, '📁')
                lines.append(f"- {emoji} **{ftype.upper()}：** {count} 个文件")
        return '\n'.join(lines) if lines else '- 无文件'
    
    @property
    def summary(self) -> Dict[str, Any]:
        """获取摘要（懒加载）"""
        if not self.analysis_results:
            return self.generate_summary()
        return {
            'total_files': len(self.analysis_results),
            'file_count_by_type': {},
            'file_details': self.analysis_results
        }


def analyze_folder(folder_path: str, output_path: str = None) -> Dict[str, Any]:
    """便捷函数：分析文件夹"""
    analyzer = DataAnalyzer(folder_path)
    return analyzer.generate_summary()


def analyze_file(file_path: str) -> Dict[str, Any]:
    """便捷函数：分析单个文件"""
    analyzer = DataAnalyzer()
    return analyzer.analyze_file(file_path)


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) > 1:
        path = sys.argv[1]
        analyzer = DataAnalyzer(path)
        result = analyzer.generate_summary()
        print(f"分析完成: {result['total_files']} 个文件")
    else:
        print("用法: python data_analyzer.py <文件夹路径>")
